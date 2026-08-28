from pathlib import Path
import os
import re
import zipfile

from .online_sources import load_online_sources


SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx", ".xls", ".txt", ".md",
    ".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff",
}
MAX_ARCHIVE_MEMBERS = int(os.getenv("MAX_ARCHIVE_MEMBERS", "10000"))
MAX_ARCHIVE_EXPANDED_BYTES = int(os.getenv("MAX_ARCHIVE_EXPANDED_BYTES", str(200 * 1024 * 1024)))
MAX_IMAGE_PIXELS = int(os.getenv("MAX_IMAGE_PIXELS", "40000000"))
MAX_PDF_PAGES = int(os.getenv("MAX_PDF_PAGES", "5000"))

YEAR_PATTERN = re.compile(r"(?<!\d)((?:19|20)\d{2})(?:\s*年|\s*版|\s*版本)?(?!\d)")
VERSION_PATTERN = re.compile(r"(?i)(?:^|[\s_\-（(])v(?:ersion)?\s*(\d+(?:\.\d+)*)")
AMENDMENT_PATTERN = re.compile(
    r"(修订通知|修订决定|修改通知|修改决定|补充规定|补充通知|补充协议|勘误|变更通知)"
)
VERSION_DECORATION_PATTERN = re.compile(
    r"(?i)[\s_\-—]*(修订版|最新版|最终版|正式版|完整版|全文)$"
)


def get_version_info(path):
    """从文件名提取文档系列和版本；不让大模型猜测版本。"""
    path = Path(path)
    stem = path.stem.strip()
    year_match = YEAR_PATTERN.search(stem)
    version_match = VERSION_PATTERN.search(stem)
    year = int(year_match.group(1)) if year_match else None
    version = version_match.group(1) if version_match else None
    document_kind = "amendment" if AMENDMENT_PATTERN.search(stem) else "full"

    family = YEAR_PATTERN.sub("", stem)
    family = VERSION_PATTERN.sub(" ", family)
    family = AMENDMENT_PATTERN.sub("", family)
    family = VERSION_DECORATION_PATTERN.sub("", family)
    family = re.sub(r"[\s_\-—（）()\[\]【】]+", "", family).lower()
    family = family or stem.lower()

    if year is not None:
        version_key = (2, year)
        version_label = f"{year}版"
    elif version is not None:
        parts = tuple(int(part) for part in version.split("."))
        version_key = (1, *parts)
        version_label = f"V{version}"
    else:
        version_key = (0,)
        version_label = None

    return {
        "document_family": family,
        "version_year": year,
        "version_label": version_label,
        "version_key": version_key,
        "has_explicit_version": year is not None or version is not None,
        "document_kind": document_kind,
    }


def apply_version_metadata(items):
    file_info = {}
    for item in items:
        path = item["file_path"]
        if path not in file_info:
            file_info[path] = get_version_info(path)

    active_paths_by_family = _active_paths_by_family(file_info)

    for item in items:
        info = file_info[item["file_path"]]
        item.update(info)
        item["is_latest_version"] = item["file_path"] in active_paths_by_family[
            info["document_family"]
        ]
        item["effective_order"] = info["version_key"]
    return items


def _active_paths_by_family(file_info, cutoff_year=None):
    """选择每个系列的最新完整版本，以及其后的所有增量修订。"""
    grouped = {}
    for path, info in file_info.items():
        if cutoff_year is not None:
            year = info.get("version_year")
            if year is not None and year > cutoff_year:
                continue
        grouped.setdefault(info["document_family"], []).append((path, info))

    selected = {}
    for family, entries in grouped.items():
        full = [entry for entry in entries if entry[1]["document_kind"] == "full"]
        if full:
            baseline_key = max(entry[1]["version_key"] for entry in full)
            paths = {path for path, info in full if info["version_key"] == baseline_key}
            paths.update(
                path for path, info in entries
                if info["document_kind"] == "amendment"
                and info["version_key"] >= baseline_key
            )
        else:
            # 资料库只有修订文件时全部保留，交给回答阶段明确提示依据不完整。
            paths = {path for path, _ in entries}
        selected[family] = paths
    return selected


def filter_chunks_for_question(chunks, question):
    """明确问年份时重建当时有效链，否则使用当前有效链。"""
    year_match = YEAR_PATTERN.search(question)
    if year_match:
        requested_year = int(year_match.group(1))
        file_info = {}
        for item in chunks:
            file_info.setdefault(item["file_path"], {
                key: item.get(key) for key in (
                    "document_family", "version_year", "version_key", "document_kind"
                )
            })
        active = _active_paths_by_family(file_info, cutoff_year=requested_year)
        active_paths = set().union(*active.values()) if active else set()
        matching = [item for item in chunks if item["file_path"] in active_paths]
        if matching:
            return matching, f"已重建截至 {requested_year} 年有效的完整版本与后续修订"
        return [], f"知识库中没有识别到 {requested_year} 版本资料"

    latest = [item for item in chunks if item.get("is_latest_version", True)]
    return latest, "已使用最新完整版本，并叠加其后的有效修订"


def _item(path, location, text):
    return {
        "file_name": path.name,
        "file_path": str(path),
        "location": location,
        "text": text.strip(),
    }


def _load_pdf(path):
    import pymupdf

    items = []
    with pymupdf.open(path) as doc:
        if len(doc) > MAX_PDF_PAGES:
            raise RuntimeError(f"PDF 页数超过安全限制：{MAX_PDF_PAGES}")
        for page_number, page in enumerate(doc, start=1):
            text = page.get_text().strip()
            if text:
                items.append(_item(path, f"第 {page_number} 页", text))
    return items


def _load_docx(path):
    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError("读取 Word 需要安装 python-docx：pip install python-docx") from error

    document = Document(path)
    items = []

    for number, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            items.append(_item(path, f"段落 {number}", text))

    for table_number, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                rows.append(" | ".join(values))
        if rows:
            items.append(_item(path, f"表格 {table_number}", "\n".join(rows)))

    return items


def _load_xlsx(path):
    try:
        import openpyxl
    except ImportError as error:
        raise RuntimeError("读取 Excel 需要安装 openpyxl：pip install openpyxl") from error

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    items = []
    try:
        for sheet in workbook.worksheets:
            for row_number, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                values = ["" if value is None else str(value).strip() for value in row]
                if any(values):
                    text = " | ".join(values)
                    items.append(_item(path, f"工作表 {sheet.title} 第 {row_number} 行", text))
    finally:
        workbook.close()
    return items


def _load_xls(path):
    try:
        import pandas as pd
    except ImportError as error:
        raise RuntimeError("读取旧版 Excel 需要安装 pandas 和 xlrd：pip install pandas xlrd") from error

    items = []
    sheets = pd.read_excel(path, sheet_name=None, header=None)
    for sheet_name, frame in sheets.items():
        for row_number, row in enumerate(frame.itertuples(index=False, name=None), start=1):
            values = ["" if pd.isna(value) else str(value).strip() for value in row]
            if any(values):
                items.append(
                    _item(path, f"工作表 {sheet_name} 第 {row_number} 行", " | ".join(values))
                )
    return items


def _load_text(path):
    for encoding in ("utf-8-sig", "gb18030"):
        try:
            text = path.read_text(encoding=encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise RuntimeError("无法识别文本文件编码")

    return [_item(path, "全文", text)] if text.strip() else []


def _load_image(path):
    try:
        import pytesseract
        from PIL import Image
    except ImportError as error:
        raise RuntimeError(
            "读取图片文字需要安装 Pillow 和 pytesseract，并安装 Tesseract OCR："
            "pip install pillow pytesseract"
        ) from error

    with Image.open(path) as image:
        width, height = image.size
        if width * height > MAX_IMAGE_PIXELS:
            raise RuntimeError(f"图片像素超过安全限制：{MAX_IMAGE_PIXELS}")
        image.load()
        text = pytesseract.image_to_string(image, lang="chi_sim+eng")
    return [_item(path, "图片 OCR", text)] if text.strip() else []


def _validate_archive_infos(infos):
    if len(infos) > MAX_ARCHIVE_MEMBERS:
        raise RuntimeError(f"Office 文件条目数超过安全限制：{MAX_ARCHIVE_MEMBERS}")
    expanded = sum(max(0, item.file_size) for item in infos)
    if expanded > MAX_ARCHIVE_EXPANDED_BYTES:
        raise RuntimeError(f"Office 文件解压后超过安全限制：{MAX_ARCHIVE_EXPANDED_BYTES} 字节")
    if any(item.flag_bits & 0x1 for item in infos):
        raise RuntimeError("不支持加密的 Office 文件")


def _validate_archive(path):
    try:
        with zipfile.ZipFile(path) as archive:
            _validate_archive_infos(archive.infolist())
    except zipfile.BadZipFile as error:
        raise RuntimeError("Office 文件结构无效") from error


def load_file(path):
    path = Path(path)
    extension = path.suffix.lower()
    if extension in {".docx", ".xlsx"}:
        _validate_archive(path)
    loaders = {
        ".pdf": _load_pdf,
        ".docx": _load_docx,
        ".xlsx": _load_xlsx,
        ".xls": _load_xls,
        ".txt": _load_text,
        ".md": _load_text,
        ".png": _load_image,
        ".jpg": _load_image,
        ".jpeg": _load_image,
        ".bmp": _load_image,
        ".tif": _load_image,
        ".tiff": _load_image,
    }
    return loaders[extension](path)


def load_source_directory(source_directory):
    source_directory = Path(source_directory)
    files = sorted(
        path for path in source_directory.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )

    if not files:
        raise RuntimeError(f"知识库目录中没有支持的文件：{source_directory}")

    items = []
    failures = []
    for path in files:
        try:
            loaded = load_file(path)
            items.extend(loaded)
            print(f"已读取：{path.name}（{len(loaded)} 个内容单元）")
        except Exception as error:
            failures.append(f"{path.name}：{error}")
            print(f"跳过文件：{path.name}；原因：{error}")

    if not items:
        details = "\n".join(failures)
        raise RuntimeError(f"没有成功读取任何资料。\n{details}")

    return apply_version_metadata(items)


def load_all_sources(source_directory):
    """统一加载本地文件与已配置的在线文档。"""
    source_directory = Path(source_directory)
    items = []
    failures = []
    try:
        items.extend(load_source_directory(source_directory))
    except RuntimeError as error:
        failures.append(str(error))

    config_path = Path(os.getenv(
        "ONLINE_SOURCES_CONFIG",
        str(source_directory / "online_sources.json"),
    ))
    online_items, online_failures = load_online_sources(config_path)
    items.extend(online_items)
    failures.extend(online_failures)
    for failure in online_failures:
        print(f"跳过在线文档：{failure}")

    if not items:
        raise RuntimeError("没有成功读取任何资料。\n" + "\n".join(failures))
    return apply_version_metadata(items), failures
